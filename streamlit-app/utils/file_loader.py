"""
file_loader.py
HK-AICOS Phase 2.0 - File Loading Utilities

Handles uploaded files (images, PDFs, DOCX, XLSX) and extracts content for analysis.
"""

import os
import uuid
from datetime import datetime
from pathlib import Path

from PIL import Image
import pypdf

UPLOAD_DIR = Path(__file__).parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".pdf", ".docx", ".xlsx"}
MAX_FILE_SIZE_MB = 50

# Human-readable type labels (Traditional Chinese)
FILE_TYPE_LABELS = {
    ".jpg": "圖片",
    ".jpeg": "圖片",
    ".png": "圖片",
    ".pdf": "PDF",
    ".docx": "Word 文件",
    ".xlsx": "Excel 試算表",
}


def save_uploaded_file(uploaded_file) -> Path:
    """Save a Streamlit uploaded file to the uploads directory. Returns the saved path."""
    ext = Path(uploaded_file.name).suffix.lower()
    unique_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}{ext}"
    save_path = UPLOAD_DIR / unique_name
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return save_path


def is_allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def get_file_type_label(filename: str) -> str:
    """Return a human-readable file type label in Traditional Chinese."""
    ext = Path(filename).suffix.lower()
    return FILE_TYPE_LABELS.get(ext, ext.upper().lstrip("."))


def extract_pdf_text(file_path: Path) -> str:
    """Extract text content from a PDF file."""
    try:
        reader = pypdf.PdfReader(str(file_path))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {i + 1}]\n{page_text.strip()}")
        return "\n\n".join(text_parts) if text_parts else "[PDF contains no extractable text - may be scanned image]"
    except Exception as e:
        return f"[Error reading PDF: {str(e)}]"


def extract_docx_text(file_path: Path) -> str:
    """
    Extract text from a DOCX file.
    Extracts both paragraph text and table cell text.
    """
    try:
        from docx import Document  # python-docx
        doc = Document(str(file_path))
        parts = []

        # Paragraphs
        para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if para_texts:
            parts.append("【段落內容】\n" + "\n".join(para_texts))

        # Tables
        for t_idx, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats merged cell text)
                seen = []
                for c in cells:
                    if not seen or c != seen[-1]:
                        seen.append(c)
                row_line = " | ".join(seen)
                if row_line.strip():
                    rows_text.append(row_line)
            if rows_text:
                parts.append(f"【表格 {t_idx + 1}】\n" + "\n".join(rows_text))

        return "\n\n".join(parts) if parts else "[DOCX contains no extractable text]"
    except ImportError:
        return "[Error: python-docx not installed. Please add python-docx to requirements.txt]"
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


def extract_xlsx_text(file_path: Path, max_rows: int = 200) -> str:
    """
    Extract text from an XLSX file.
    Reads sheet names and up to max_rows rows per sheet.
    """
    try:
        import openpyxl  # openpyxl
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows:
                    rows_text.append(f"[... 已達 {max_rows} 行上限，餘下內容略去 ...]")
                    break
                # Skip completely empty rows
                cell_values = [str(c).strip() if c is not None else "" for c in row]
                if any(cell_values):
                    rows_text.append(" | ".join(cell_values))
                    row_count += 1

            if rows_text:
                parts.append(f"【Sheet：{sheet_name}】\n" + "\n".join(rows_text))
            else:
                parts.append(f"【Sheet：{sheet_name}】（空白）")

        wb.close()
        return "\n\n".join(parts) if parts else "[XLSX contains no extractable data]"
    except ImportError:
        return "[Error: openpyxl not installed. Please add openpyxl to requirements.txt]"
    except Exception as e:
        return f"[Error reading XLSX: {str(e)}]"


def get_image_description(file_path: Path) -> str:
    """Get basic image metadata as description."""
    try:
        with Image.open(file_path) as img:
            width, height = img.size
            mode = img.mode
            fmt = img.format or Path(file_path).suffix.upper().lstrip(".")
            return f"Image file: {fmt}, {width}x{height}px, mode={mode}"
    except Exception as e:
        return f"[Error reading image: {str(e)}]"


def load_file_content(file_path: Path) -> dict:
    """
    Load file and return a dict with type, description, and raw content.

    Supported types: image (jpg/jpeg/png), pdf, docx, xlsx.

    Returns:
        {
            "type": "image" | "pdf" | "docx" | "xlsx" | "unknown",
            "filename": str,
            "description": str,   # short description for prompt
            "content": str,        # extracted text or metadata
            "path": Path,
        }
    """
    ext = file_path.suffix.lower()
    filename = file_path.name

    if ext in {".jpg", ".jpeg", ".png"}:
        description = get_image_description(file_path)
        return {
            "type": "image",
            "filename": filename,
            "description": description,
            "content": description,
            "path": file_path,
        }

    elif ext == ".pdf":
        text = extract_pdf_text(file_path)
        if len(text) > 8000:
            text = text[:8000] + "\n\n[... content truncated for analysis ...]"
        description = f"PDF 文件（已抽取 {len(text)} 字元）"
        return {
            "type": "pdf",
            "filename": filename,
            "description": description,
            "content": text,
            "path": file_path,
        }

    elif ext == ".docx":
        text = extract_docx_text(file_path)
        if len(text) > 8000:
            text = text[:8000] + "\n\n[... content truncated for analysis ...]"
        description = f"Word 文件（已抽取 {len(text)} 字元）"
        return {
            "type": "docx",
            "filename": filename,
            "description": description,
            "content": text,
            "path": file_path,
        }

    elif ext == ".xlsx":
        text = extract_xlsx_text(file_path, max_rows=200)
        if len(text) > 8000:
            text = text[:8000] + "\n\n[... content truncated for analysis ...]"
        description = f"Excel 試算表（已抽取 {len(text)} 字元）"
        return {
            "type": "xlsx",
            "filename": filename,
            "description": description,
            "content": text,
            "path": file_path,
        }

    else:
        return {
            "type": "unknown",
            "filename": filename,
            "description": "不支援的文件格式",
            "content": "",
            "path": file_path,
        }


def process_uploaded_file(uploaded_file) -> dict:
    """
    Full pipeline: save uploaded file then load its content.
    Returns content dict (see load_file_content).
    """
    saved_path = save_uploaded_file(uploaded_file)
    return load_file_content(saved_path)
